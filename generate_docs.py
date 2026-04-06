from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_cahier_des_charges():
    doc = Document()
    title = doc.add_heading("Cahier des Charges - Elyon 360", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading("1. Présentation du Projet", level=1)
    doc.add_paragraph("ElyonSys 360 est une plateforme SaaS (Software as a Service) multi-tenant conçue pour révolutionner la gestion des églises. Elle offre une solution intégrée pour l'administration interne, l'engagement des membres et la visibilité publique.")

    doc.add_heading("2. Objectifs", level=1)
    doc.add_paragraph("- Centraliser la gestion administrative (membres, finances, logistique).")
    doc.add_paragraph("- Offrir un site vitrine professionnel à chaque église.")
    doc.add_paragraph("- Faciliter l'interaction entre l'église et ses membres via un espace dédié.")
    doc.add_paragraph("- Automatiser les processus récurrents comme les rapports de dons et la gestion des cérémonies.")

    doc.add_heading("3. Public Cible", level=1)
    doc.add_paragraph("- Églises de toutes tailles (locales ou réseaux).")
    doc.add_paragraph("- Responsables d'églises (Pasteurs, Administrateurs, Trésoriers).")
    doc.add_paragraph("- Membres et fidèles de l'église.")

    doc.add_heading("4. Périmètre Fonctionnel", level=1)
    doc.add_heading("4.1 Portail SaaS (ElyonSys)", level=2)
    doc.add_paragraph("- Inscription et onboarding des églises.")
    doc.add_paragraph("- Gestion des abonnements et paiements.")
    doc.add_paragraph("- Administration globale du système.")
    
    doc.add_heading("4.2 Espace Administratif (Back-office)", level=2)
    doc.add_paragraph("- Gestion des membres (Profils, types, catégories).")
    doc.add_paragraph("- Gestion des groupes, ministères et départements.")
    doc.add_paragraph("- Comptabilité (Gestion des dons, dîmes, offrandes et promesses).")
    doc.add_paragraph("- École du Dimanche (EDD) : Suivi des cours et présences.")
    doc.add_paragraph("- Gestion de la logistique et de l'inventaire.")
    doc.add_paragraph("- Gestion des cérémonies (Mariages, Baptêmes, Sainte Cène).")

    doc.add_heading("4.3 Espace Membre", level=2)
    doc.add_paragraph("- Authentification personnelle.")
    doc.add_paragraph("- Consultation du profil et historique de participation.")
    doc.add_paragraph("- Suivi personnel des dons et contributions.")
    doc.add_paragraph("- Réception de notifications et messages de l'administration.")

    doc.add_heading("4.4 Site Public (CMS)", level=2)
    doc.add_paragraph("- Présentation de l'église et de l'équipe pastorale.")
    doc.add_paragraph("- Calendrier des événements et horaires de culte.")
    doc.add_paragraph("- Formulaire de contact et inscription pour nouveaux visiteurs.")

    doc.add_heading("5. Contraintes", level=1)
    doc.add_paragraph("- Sécurité : Isolation stricte des données par église.")
    doc.add_paragraph("- Disponibilité : Service accessible 24/7.")
    doc.add_paragraph("- Ergonomie : Interface intuitive et moderne adaptée à divers profils d'utilisateurs.")

    doc.save("Cahier_des_Charges_Elyon360.docx")

def create_specification_technique():
    doc = Document()
    title = doc.add_heading("Spécifications Techniques - Elyon 360", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Architecture Globale", level=1)
    doc.add_paragraph("L'application repose sur une architecture découplée (Client-Serveur) favorisant la scalabilité et la maintenance.")
    
    doc.add_heading("2. Stack Technologique", level=1)
    doc.add_paragraph("- Frontend : ReactJS avec Tailwind CSS pour un design responsive et moderne.")
    doc.add_paragraph("- Backend : Node.js avec le framework Express, assurant des performances élevées.")
    doc.add_paragraph("- Base de données : PostgreSQL, utilisant un modèle multi-tenant pour l'isolation des données.")
    doc.add_paragraph("- Authentification : JSON Web Tokens (JWT) avec gestion des rôles (RBAC).")
    doc.add_paragraph("- Paiements : Intégration prévue avec MonCash et d'autres passerelles.")

    doc.add_heading("3. Structure Multitenant", level=1)
    doc.add_paragraph("Le système utilise des sous-domaines dynamiques (ex: eglise.elyonsys360.com). Chaque requête backend identifie l'église concernée via le contexte du sous-domaine ou des headers spécifiques, garantissant que les données d'une église ne fuitent jamais vers une autre.")

    doc.add_heading("4. Modules Clés", level=1)
    doc.add_heading("4.1 Moteur de Comptabilité", level=2)
    doc.add_paragraph("Gestion des transactions financières avec support multi-devises et rapports statistiques dynamiques.")
    
    doc.add_heading("4.2 Système de Gestion de Contenu (CMS)", level=2)
    doc.add_paragraph("Permet aux administrateurs de l'église de modifier les sections du site public sans compétences techniques.")

    doc.add_heading("5. Sécurité et Infrastructure", level=1)
    doc.add_paragraph("- Chiffrement des mots de passe avec bcrypt.")
    doc.add_paragraph("- Protection contre les injections SQL via l'utilisation d'un ORM (Sequelize/Prisma).")
    doc.add_paragraph("- Déploiement sur infrastructure Cloud (AWS/Heroku/Vercel).")

    doc.save("Specification_Technique_Elyon360.docx")

def create_manuel_utilisation():
    doc = Document()
    title = doc.add_heading("Manuel d'Utilisation - Elyon 360", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph("Ce manuel guide les utilisateurs dans l'utilisation quotidienne de la plateforme Elyon 360.")

    doc.add_heading("2. Pour les Administrateurs", level=1)
    doc.add_heading("2.1 Tableau de Bord", level=2)
    doc.add_paragraph("Dès la connexion, vous accédez à une vue d'ensemble : nombre de membres actifs, total des dons du mois et événements à venir.")
    
    doc.add_heading("2.2 Gestion des Membres", level=2)
    doc.add_paragraph("Pour ajouter un membre : allez dans 'Membres', cliquez sur 'Ajouter', remplissez les informations et assignez-le à un type ou département.")
    
    doc.add_heading("2.3 Comptabilité", level=2)
    doc.add_paragraph("Enregistrez chaque don reçu pour assurer la transparence. Utilisez la section 'Rapports' pour exporter vos bilans financiers.")

    doc.add_heading("3. Pour les Membres", level=1)
    doc.add_paragraph("Connectez-vous via l'URL de votre église. Dans votre espace personnel, vous pouvez mettre à jour vos coordonnées, consulter vos engagements et vos dons.")

    doc.add_heading("4. Pour le Super Admin", level=1)
    doc.add_paragraph("Le portail ElyonSys permet de valider les nouvelles inscriptions d'églises, fixer les limites des abonnements et monitorer la santé globale du SaaS.")

    doc.save("Manuel_Utilisation_Elyon360.docx")

if __name__ == "__main__":
    create_cahier_des_charges()
    create_specification_technique()
    create_manuel_utilisation()
    print("Documents générés avec succès.")
